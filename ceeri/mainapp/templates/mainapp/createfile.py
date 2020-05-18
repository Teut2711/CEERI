
import os
from docx import Document
from itertools import zip_longest

def main():

    doc = Document("CEERI.docx")

    doc.tables[1].rows[0].cells[1].text = form1

    doc.tables[2].rows[0].cells[1].text = form2

    for i, j in zip_longest(form3,[
        cell
        for x in doc.tables[3].rows[1:]
        for cell in x.cells
    ], fillvalue=""):
        j.cells.text = u"\u25A1"+" "+i+" "
         

    doc.tables[4].rows[0].cells[1].text = form4

    def form5_process(): 
        doc.tables[5].rows[0].cells[0].text  += form5.title 
        doc.tables[5].rows[1].cells[0].text += form5.name
        doc.tables[5].rows[1].cells[1].text = forms.sex

        for i, j in zip(doc.tables[5].rows, form5):  
            i.cells[1].text = forms5.fulladress 
            i.cells[1].text = forms5.mobile 
            i.cells[1].text = forms5.fax 
            i.cells[1].text = forms5.email  

    def form6_process():
        doc.tables[6].rows [0].cells[0].text += form6
        doc.tables[6].rows [0].cells[1].text += form6
        
        