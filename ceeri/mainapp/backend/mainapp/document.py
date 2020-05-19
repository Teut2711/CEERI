import os
from docx import Document
from itertools import zip_longest, tee
from more_itertools import roundrobin

from mainapp.forms import *
from more_itertools.recipes import partition
from operator import itemgetter
import datetime

doc = None


def form1_process(form:Form1):
    
    doc.tables[1].rows[0].cells[1].text  = form.cleaned_data["institute"]


def form2_process(form:Form2):
    
    doc.tables[2].rows[0].cells[1].text = form.cleaned_data["project_title"]
    


def form3_process(form:Form3):


    ## correct order of checkbox 
    all_choices = form.fields["subjects"]._choices
    for k, i in enumerate(all_choices):
        all_choices[k] = [    i[1], 
                          i[0] in form.cleaned_data["subjects"]] 
         
    
    for i, j in zip((*roundrobin(all_choices[:5], all_choices[5:8], all_choices[8:]),), [
        cell
        for x in doc.tables[3].rows[1:]
        for cell in x.cells
    ]):
       
        if i[1]:
            j.text = u"\u2611"+" "+i[0]+" "
        else:
            j.text = u"\u2610"+" "+i[0]+" "
    
            

def form4_process(form:Form4):
    doc.tables[4].rows[0].cells[1].text = ""
    for i, j in form.cleaned_data.items():
        if j:
            doc.tables[4].rows[0].cells[1].text += i.title()+" : " +j +"\n"
             

def form5_process(form:Form5):
    cleaned_data = list(form.cleaned_data.keys())
    doc.tables[5].rows[0].cells[1].text = form.cleaned_data['title'].capitalize()
    doc.tables[5].rows[1].cells[1].text = form.cleaned_data['name'].capitalize()
    doc.tables[5].rows[1].cells[2].text = form.cleaned_data['sex'].capitalize()

    cleaned_data = cleaned_data[3:]

    for i, j in zip(doc.tables[5].rows[2:], cleaned_data):
        i.cells[1].text =str(form.cleaned_data[j]).capitalize()
   

def form6_process(form):
    pass
    doc.tables[6].rows[0].cells[0].text += form6
    doc.tables[6].rows[0].cells[1].text += form6



def main(forms_dict ):
    
    global doc   
    
    doc = Document(os.path.join(
        os.path.dirname(__file__),
        "CEERI.docx"))
    for k, i in enumerate(forms_dict.values()):
        if k+1<= 5:
            exec(f"form{k+1}_process(i)")
    
    doc.save(r"C:\Users\Dell\Desktop\github\ceeri\ceeri\mainapp\backend\mainapp\junk.docx")
